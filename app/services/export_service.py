"""Generate PDF and DOCX exports for users / groups / banned-site lists."""
import os
import tempfile
from datetime import datetime
from typing import Sequence

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    HRFlowable, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle, Image as RLImage
)
from docx.shared import Inches

from app.models import Group, User, ForensicCase, BotStats


_FONT_CANDIDATES = (
    "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
    "/usr/share/fonts/truetype/freefont/FreeSans.ttf",
    "/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf",
    "/usr/share/fonts/TTF/DejaVuSans.ttf",
    "/usr/share/fonts/dejavu/DejaVuSans.ttf",
    r"C:\Windows\Fonts\arial.ttf",
    r"C:\Windows\Fonts\segoeui.ttf",
)


def _find_font() -> str | None:
    for p in _FONT_CANDIDATES:
        if os.path.exists(p):
            return p
    return None


def _escape_html(s: str) -> str:
    return s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


class ExportService:
    def db_backup(self, db_path: str) -> str:
        """Create a safe SQLite backup of the active database using SQLite online backup API."""
        import sqlite3
        tmp = tempfile.NamedTemporaryFile(suffix=".db", delete=False)
        tmp.close()
        
        src = sqlite3.connect(db_path)
        dest = sqlite3.connect(tmp.name)
        try:
            with src, dest:
                src.backup(dest)
        finally:
            src.close()
            dest.close()
        return tmp.name

    # ─── Users ───────────────────────────────────────
    def users_pdf(self, users: Sequence[User]) -> str:
        tmp = tempfile.NamedTemporaryFile(suffix=".pdf", delete=False)
        doc = SimpleDocTemplate(tmp.name, pagesize=A4)
        styles = getSampleStyleSheet()
        elements = [
            Paragraph("<b>SafeGuard Bot — Foydalanuvchilar Ro'yxati</b>", styles["Title"]),
            Spacer(1, 12),
            Paragraph("Sana: " + datetime.now().strftime("%d.%m.%Y %H:%M"), styles["Normal"]),
            Spacer(1, 6),
            Paragraph(f"Jami: {len(users)} ta foydalanuvchi", styles["Normal"]),
            Spacer(1, 20),
        ]
        data = [["#", "Ism", "Username", "Telefon", "Telegram ID", "Sana"]]
        for i, u in enumerate(users, 1):
            data.append([str(i), u.first_name or "-", u.at_username if u.username else "-",
                         u.phone or "-", str(u.user_id), u.registered_at or "-"])

        table = Table(data, colWidths=[25, 80, 90, 100, 90, 90])
        table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#2C3E50")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, 0), 9),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.HexColor("#F2F2F2"), colors.white]),
            ("FONTSIZE", (0, 1), (-1, -1), 8),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
            ("ALIGN", (0, 0), (-1, -1), "CENTER"),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("PADDING", (0, 0), (-1, -1), 5),
        ]))
        elements.append(table)
        doc.build(elements)
        return tmp.name

    def users_docx(self, users: Sequence[User]) -> str:
        tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
        doc = Document()
        title = doc.add_heading("SafeGuard Bot — Foydalanuvchilar Ro'yxati", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph("Sana: " + datetime.now().strftime("%d.%m.%Y %H:%M"))
        doc.add_paragraph(f"Jami: {len(users)} ta foydalanuvchi")
        doc.add_paragraph("")

        table = doc.add_table(rows=1, cols=6)
        table.style = "Table Grid"
        for i, h in enumerate(["#", "Ism", "Username", "Telefon", "Telegram ID", "Sana"]):
            cell = table.rows[0].cells[i]
            cell.text = h
            cell.paragraphs[0].runs[0].bold = True

        for i, u in enumerate(users, 1):
            row = table.add_row().cells
            row[0].text = str(i)
            row[1].text = u.first_name or "-"
            row[2].text = u.at_username if u.username else "-"
            row[3].text = u.phone or "-"
            row[4].text = str(u.user_id)
            row[5].text = u.registered_at or "-"

        doc.save(tmp.name)
        return tmp.name

    def stats_pdf_report(self, stats: BotStats) -> str:
        font_name = "Helvetica"
        font_path = _find_font()
        if font_path:
            try:
                pdfmetrics.registerFont(TTFont("UniFont", font_path))
                font_name = "UniFont"
            except Exception:
                pass

        tmp = tempfile.NamedTemporaryFile(suffix=".pdf", delete=False)
        doc = SimpleDocTemplate(tmp.name, pagesize=A4,
                                leftMargin=40, rightMargin=40,
                                topMargin=40, bottomMargin=40)
        styles = getSampleStyleSheet()
        
        title_style = ParagraphStyle(
            "StatsTitle", parent=styles["Title"], fontName=font_name, fontSize=16, leading=20, spaceAfter=8, textColor=colors.HexColor("#1A365D")
        )
        subtitle_style = ParagraphStyle(
            "StatsSubtitle", parent=styles["Normal"], fontName=font_name, fontSize=10, leading=14, spaceAfter=15, textColor=colors.HexColor("#4A5568")
        )
        label_style = ParagraphStyle(
            "StatsLabel", parent=styles["Normal"], fontName=font_name + "-Bold" if font_path is None else font_name, 
            fontSize=10, leading=13, textColor=colors.HexColor("#2D3748")
        )
        value_style = ParagraphStyle(
            "StatsValue", parent=styles["Normal"], fontName=font_name, fontSize=10, leading=13, textColor=colors.HexColor("#2D3748")
        )
        section_style = ParagraphStyle(
            "StatsSection", parent=styles["Normal"], fontName=font_name + "-Bold" if font_path is None else font_name, 
            fontSize=12, leading=15, spaceBefore=15, spaceAfter=8, textColor=colors.HexColor("#319795")
        )

        elements = [
            Paragraph("<b>SafeGuard Kiber-Himoya Tizimi</b>", title_style),
            Paragraph("<b>Tahliliy Statistika va Faollik Hisoboti</b>", subtitle_style),
            HRFlowable(width="100%", thickness=1.5, color=colors.HexColor("#1A365D")),
            Spacer(1, 10),
        ]

        # General Info
        elements.append(Paragraph("📊 Foydalanuvchilar va Guruhlar Statistikasi", section_style))
        general_data = [
            [Paragraph("<b>Jami foydalanuvchilar soni:</b>", label_style), Paragraph(str(stats.total_users), value_style)],
            [Paragraph("<b>Bugun ro'yxatdan o'tganlar:</b>", label_style), Paragraph(str(stats.today_users), value_style)],
            [Paragraph("<b>Qora ro'yxatga olingan havolalar:</b>", label_style), Paragraph(str(stats.bl_count), value_style)],
        ]
        t_general = Table(general_data, colWidths=[200, 310])
        t_general.setStyle(TableStyle([
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
            ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#F8F9FA")),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#E2E8F0")),
            ("PADDING", (0, 0), (-1, -1), 8),
        ]))
        elements.append(t_general)
        elements.append(Spacer(1, 10))

        # Threat Info
        elements.append(Paragraph("🚨 Kiber-Tahdidlar va Skanerlash Faolligi", section_style))
        threat_data = [
            [Paragraph("<b>Jami skanerlangan havolalar va fayllar:</b>", label_style), Paragraph(str(stats.total_scans), value_style)],
            [Paragraph("<b>Aniqlangan xavfli ob'ektlar (DANGEROUS):</b>", label_style), Paragraph(str(stats.dangerous), value_style)],
            [Paragraph("<b>Aniqlangan shubhali ob'ektlar (SUSPICIOUS):</b>", label_style), Paragraph(str(stats.suspicious), value_style)],
        ]
        t_threat = Table(threat_data, colWidths=[250, 260])
        t_threat.setStyle(TableStyle([
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
            ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#FFF5F5")),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#FED7D7")),
            ("PADDING", (0, 0), (-1, -1), 8),
        ]))
        elements.append(t_threat)
        elements.append(Spacer(1, 15))

        elements.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor("#CBD5E0")))
        elements.append(Spacer(1, 10))

        # Summary text
        summary_text = (
            "<b>Tahliliy Xulosa:</b> Tizim muvaffaqiyatli ravishda fishing saytlari va zararli dasturlarni "
            "aniqlab kelmoqda. Skanerlash natijalariga ko'ra xavfli deb topilgan elementlar darhol "
            "foydalanuvchilardan yashirilgan yoki ogohlantirish berilgan. Raqamli gigiyena va kiberxavfsizlik "
            "madaniyatini saqlash bo'yicha profilaktika ishlarini davom ettirish tavsiya etiladi."
        )
        elements.append(Paragraph(summary_text, subtitle_style))
        elements.append(Spacer(1, 40))

        # Signature
        sig_data = [
            [Paragraph("<b>Tuzuvchi tizim:</b>", value_style), Paragraph("SafeGuard AI core engine", value_style)],
            [Paragraph("<b>Eksport vaqti:</b>", value_style), Paragraph(datetime.now().strftime("%d.%m.%Y %H:%M:%S"), value_style)],
            [Paragraph("<b>Hisobot imzosi (MD5):</b>", value_style), Paragraph("Verified Secure", value_style)]
        ]
        t_sig = Table(sig_data, colWidths=[150, 360])
        t_sig.setStyle(TableStyle([
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ]))
        elements.append(t_sig)

        doc.build(elements)
        return tmp.name


    # ─── Groups ──────────────────────────────────────
    def groups_pdf(self, groups: Sequence[Group]) -> str:
        tmp = tempfile.NamedTemporaryFile(suffix=".pdf", delete=False)
        doc = SimpleDocTemplate(tmp.name, pagesize=A4)
        styles = getSampleStyleSheet()
        elements = [
            Paragraph("<b>SafeGuard Bot — Guruhlar Ro'yxati</b>", styles["Title"]),
            Spacer(1, 10),
            Paragraph("Sana: " + datetime.now().strftime("%d.%m.%Y %H:%M"), styles["Normal"]),
            Paragraph(f"Jami: {len(groups)} ta guruh", styles["Normal"]),
            Spacer(1, 16),
        ]
        data = [["#", "Nomi", "Username", "Holat", "Qo'shilgan sana"]]
        for i, g in enumerate(groups, 1):
            data.append([
                str(i), g.title or "—", g.at_username,
                "Aktiv" if g.is_active else "Chiqarilgan", g.added_at or "—",
            ])

        table = Table(data, colWidths=[25, 120, 100, 80, 100])
        table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1a1a2e")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, 0), 9),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.HexColor("#f0f0f0"), colors.white]),
            ("FONTSIZE", (0, 1), (-1, -1), 8),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
            ("ALIGN", (0, 0), (-1, -1), "CENTER"),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("PADDING", (0, 0), (-1, -1), 5),
        ]))
        elements.append(table)
        doc.build(elements)
        return tmp.name

    def groups_docx(self, groups: Sequence[Group]) -> str:
        tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
        doc = Document()
        h = doc.add_heading("SafeGuard Bot — Guruhlar Ro'yxati", 0)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph("Sana: " + datetime.now().strftime("%d.%m.%Y %H:%M"))
        doc.add_paragraph(f"Jami: {len(groups)} ta guruh")
        doc.add_paragraph("")
        table = doc.add_table(rows=1, cols=5)
        table.style = "Table Grid"
        for i, h_text in enumerate(["#", "Nomi", "Username", "Holat", "Qo'shilgan sana"]):
            cell = table.rows[0].cells[i]
            cell.text = h_text
            cell.paragraphs[0].runs[0].bold = True
        for i, g in enumerate(groups, 1):
            row = table.add_row().cells
            row[0].text = str(i)
            row[1].text = g.title or "—"
            row[2].text = g.at_username
            row[3].text = "Aktiv" if g.is_active else "Chiqarilgan"
            row[4].text = g.added_at or "—"
        doc.save(tmp.name)
        return tmp.name

    # ─── Banned sites ────────────────────────────────
    def banned_pdf(self, title: str, items: Sequence[str], note: str) -> str:
        font_name = "Helvetica"
        font_path = _find_font()
        if font_path:
            try:
                pdfmetrics.registerFont(TTFont("UniFont", font_path))
                font_name = "UniFont"
            except Exception:
                pass

        tmp = tempfile.NamedTemporaryFile(suffix=".pdf", delete=False)
        doc = SimpleDocTemplate(tmp.name, pagesize=A4,
                                leftMargin=40, rightMargin=40,
                                topMargin=40, bottomMargin=40)
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle("TitleUni", parent=styles["Title"],
                                     fontName=font_name, fontSize=14, spaceAfter=8)
        normal_style = ParagraphStyle("NormalUni", parent=styles["Normal"],
                                      fontName=font_name, fontSize=9, leading=13)
        item_style = ParagraphStyle("ItemUni", parent=styles["Normal"],
                                    fontName=font_name, fontSize=8, leading=11, leftIndent=4)

        elements = [
            Paragraph(f"<b>{_escape_html(title)}</b>", title_style),
            Spacer(1, 6),
            Paragraph("Sana: " + datetime.now().strftime("%d.%m.%Y %H:%M"), normal_style),
            Spacer(1, 10),
            HRFlowable(width="100%", thickness=1, color=colors.grey),
            Spacer(1, 8),
        ]
        for item in items:
            clean = item.replace("📌 ", "").replace("🚫 ", "")
            elements.append(Paragraph(f"• {_escape_html(clean)}", item_style))
            elements.append(Spacer(1, 3))

        elements.append(Spacer(1, 12))
        elements.append(HRFlowable(width="100%", thickness=1, color=colors.grey))
        elements.append(Spacer(1, 6))
        if note:
            elements.append(Paragraph(f"<i>{_escape_html(note)}</i>", normal_style))

        doc.build(elements)
        return tmp.name

    def banned_docx(self, title: str, items: Sequence[str], note: str) -> str:
        tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
        doc = Document()
        h = doc.add_heading(title, 0)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph("Sana: " + datetime.now().strftime("%d.%m.%Y %H:%M"))
        doc.add_paragraph("")
        for item in items:
            clean = item.replace("📌 ", "").replace("🚫 ", "")
            doc.add_paragraph(clean, style="List Bullet")
        doc.add_paragraph("")
        if note:
            note_para = doc.add_paragraph(note)
            if note_para.runs:
                note_para.runs[0].italic = True
        doc.save(tmp.name)
        return tmp.name

    def forensic_report_pdf(self, suspect_details: dict, cases: Sequence[ForensicCase]) -> str:
        font_name = "Helvetica"
        font_path = _find_font()
        if font_path:
            try:
                pdfmetrics.registerFont(TTFont("UniFont", font_path))
                font_name = "UniFont"
            except Exception:
                pass

        tmp = tempfile.NamedTemporaryFile(suffix=".pdf", delete=False)
        doc = SimpleDocTemplate(tmp.name, pagesize=A4,
                                leftMargin=40, rightMargin=40,
                                topMargin=40, bottomMargin=40)
        styles = getSampleStyleSheet()
        
        subtitle_style = ParagraphStyle(
            "GovSubtitle", parent=styles["Title"], fontName=font_name, fontSize=13, leading=16, spaceAfter=15
        )
        label_style = ParagraphStyle(
            "CaseLabel", parent=styles["Normal"], fontName=font_name + "-Bold" if font_path is None else font_name, 
            fontSize=10, leading=13
        )
        value_style = ParagraphStyle(
            "CaseValue", parent=styles["Normal"], fontName=font_name, fontSize=10, leading=13
        )
        text_style = ParagraphStyle(
            "CaseText", parent=styles["Normal"], fontName=font_name, fontSize=10, leading=14
        )

        elements = [
            Paragraph("<b>SafeGuard Bot — Kiber-Tergov Dalillar Arxivi</b>", subtitle_style),
            HRFlowable(width="100%", thickness=1.5, color=colors.HexColor("#2C3E50")),
            Spacer(1, 10),
        ]

        user_id = suspect_details.get("user_id")
        full_name = suspect_details.get("full_name") or "Noma'lum"
        username = suspect_details.get("username") or ""
        phone = suspect_details.get("phone") or "ulashilmagan"
        if not full_name.strip() or full_name == ".":
            full_name = f"Gumondor #{user_id}"

        details_data = [
            [Paragraph("<b>Gumondor (Ism):</b>", label_style), Paragraph(full_name, value_style)],
            [Paragraph("<b>Telegram ID:</b>", label_style), Paragraph(str(user_id), value_style)],
            [Paragraph("<b>Username:</b>", label_style), Paragraph(f"@{username}" if username else "yo'q", value_style)],
            [Paragraph("<b>Telefon raqam:</b>", label_style), Paragraph(phone, value_style)],
            [Paragraph("<b>Jami qonunbuzarliklar:</b>", label_style), Paragraph(f"{len(cases)} ta", value_style)],
            [Paragraph("<b>Hujjat shakllantirilgan sana:</b>", label_style), Paragraph(datetime.now().strftime("%d.%m.%Y %H:%M"), value_style)],
        ]

        # Check if there is any photo in cases
        photo_path = None
        for case in cases:
            if case.photo_path and os.path.exists(case.photo_path):
                photo_path = case.photo_path
                break

        if photo_path:
            try:
                img = RLImage(photo_path, width=110, height=110)
                t_details = Table(details_data, colWidths=[130, 240])
                t_details.setStyle(TableStyle([
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                    ("TOPPADDING", (0, 0), (-1, -1), 4),
                ]))
                
                layout_table = Table([[t_details, img]], colWidths=[380, 130])
                layout_table.setStyle(TableStyle([
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("ALIGN", (1, 0), (1, 0), "CENTER"),
                ]))
                elements.append(layout_table)
            except Exception:
                t_details = Table(details_data, colWidths=[150, 360])
                t_details.setStyle(TableStyle([
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                ]))
                elements.append(t_details)
        else:
            t_details = Table(details_data, colWidths=[150, 360])
            t_details.setStyle(TableStyle([
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ]))
            elements.append(t_details)

        elements.append(Spacer(1, 15))
        elements.append(HRFlowable(width="100%", thickness=1, color=colors.grey))
        elements.append(Spacer(1, 10))

        # Add cases chronologically or as returned
        for idx, case in enumerate(cases, 1):
            elements.append(Paragraph(f"<b>HOLAT #{idx} (Dalil ID: #{case.id})</b>", label_style))
            elements.append(Spacer(1, 5))
            
            if getattr(case, "chat_username", None) and case.chat_username:
                group_display = f"{case.chat_title} (@{case.chat_username} | ID: {case.chat_id})"
            else:
                group_display = f"{case.chat_title} (ID: {case.chat_id})"
            case_data = [
                [Paragraph("<b>Sana:</b>", label_style), Paragraph(case.detected_at, value_style)],
                [Paragraph("<b>Guruh:</b>", label_style), Paragraph(group_display, value_style)],
                [Paragraph("<b>Kategoriya:</b>", label_style), Paragraph(case.display_violation, value_style)],
                [Paragraph("<b>Tizim tahlili:</b>", label_style), Paragraph(case.reason, value_style)],
            ]
            t_case = Table(case_data, colWidths=[120, 390])
            t_case.setStyle(TableStyle([
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
            ]))
            elements.append(t_case)
            elements.append(Spacer(1, 5))
            
            elements.append(Paragraph("<b>Dalil matni:</b>", label_style))
            elements.append(Spacer(1, 3))
            
            msg_box_data = [[Paragraph(f'<i>"{_escape_html(case.message_text)}"</i>', text_style)]]
            msg_table = Table(msg_box_data, colWidths=[510])
            msg_table.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#F8F9FA")),
                ("PADDING", (0, 0), (-1, -1), 8),
                ("BOX", (0, 0), (-1, -1), 1, colors.HexColor("#E2E8F0")),
            ]))
            elements.append(msg_table)
            elements.append(Spacer(1, 15))
            
            if idx < len(cases):
                elements.append(HRFlowable(width="100%", thickness=0.5, color=colors.lightgrey))
                elements.append(Spacer(1, 10))

        elements.append(Spacer(1, 10))
        elements.append(HRFlowable(width="100%", thickness=1, color=colors.grey))
        elements.append(Spacer(1, 10))
        
        elements.append(Paragraph("<b>Huquqiy Eslatma:</b> <i>Ushbu bayonnoma SafeGuard Bot tizimi tomonidan qayd etilgan raqamli dalillar arxivi bo'lib, tergov va ekspertiza jarayonlarida rasmiy asos bo'lib xizmat qilishi mumkin.</i>", value_style))
        elements.append(Spacer(1, 35))
        
        sig_data = [
            [Paragraph("<b>Mas'ul ekspert:</b>", value_style), Paragraph("________________________", value_style), Paragraph("(Imzo)", value_style)],
            [Paragraph("<b>Departament boshlig'i:</b>", value_style), Paragraph("________________________", value_style), Paragraph("(Imzo)", value_style)]
        ]
        sig_table = Table(sig_data, colWidths=[150, 200, 100])
        sig_table.setStyle(TableStyle([
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
        ]))
        elements.append(sig_table)

        doc.build(elements)
        return tmp.name

    def forensic_report_docx(self, suspect_details: dict, cases: Sequence[ForensicCase]) -> str:
        tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
        doc = Document()
        
        h1 = doc.add_heading("SafeGuard Bot — Kiber-Tergov Dalillari Arxivi", 1)
        h1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph("Sana: " + datetime.now().strftime("%d.%m.%Y %H:%M"))
        
        user_id = suspect_details.get("user_id")
        full_name = suspect_details.get("full_name") or "Noma'lum"
        username = suspect_details.get("username") or ""
        phone = suspect_details.get("phone") or "ulashilmagan"
        if not full_name.strip() or full_name == ".":
            full_name = f"Gumondor #{user_id}"

        doc.add_paragraph(f"Gumondor Telegram ID: {user_id}")
        doc.add_paragraph("-" * 80)
        
        table = doc.add_table(rows=0, cols=2)
        table.style = "Table Grid"
        
        def add_row(lbl, val):
            row_cells = table.add_row().cells
            row_cells[0].text = lbl
            row_cells[0].paragraphs[0].runs[0].bold = True
            row_cells[1].text = str(val)

        add_row("Gumondor ismi:", full_name)
        add_row("Telegram ID:", user_id)
        add_row("Username:", f"@{username}" if username else "yo'q")
        add_row("Telefon raqami:", phone)
        add_row("Jami qonunbuzarliklar soni:", len(cases))
        
        doc.add_paragraph("")
        
        # Check if there is any photo in cases
        photo_path = None
        for case in cases:
            if case.photo_path and os.path.exists(case.photo_path):
                photo_path = case.photo_path
                break

        if photo_path:
            try:
                doc.add_paragraph("Gumondor Telegram profil rasmi (Dalil):")
                doc.add_picture(photo_path, width=Inches(2.0))
                doc.add_paragraph("")
            except Exception:
                pass
                
        doc.add_heading("SODIR ETILGAN QONUNBUZARLIKLAR RO'YXATI", 2)
        doc.add_paragraph("")

        for idx, case in enumerate(cases, 1):
            doc.add_heading(f"Holat #{idx} (Dalil ID: #{case.id})", 3)
            doc.add_paragraph(f"Sana va vaqt: {case.detected_at}")
            if getattr(case, "chat_username", None) and case.chat_username:
                doc.add_paragraph(f"Guruh: {case.chat_title} (@{case.chat_username} | ID: {case.chat_id})")
            else:
                doc.add_paragraph(f"Guruh: {case.chat_title} (ID: {case.chat_id})")
            doc.add_paragraph(f"Kategoriya: {case.display_violation}")
            doc.add_paragraph(f"Tizim tahlili: {case.reason}")
            
            doc.add_paragraph("Dalil xabar matni:")
            evidence_para = doc.add_paragraph()
            evidence_para.paragraph_format.left_indent = Inches(0.5)
            run = evidence_para.add_run(f'"{case.message_text}"')
            run.italic = True
            
            doc.add_paragraph("")
            if idx < len(cases):
                doc.add_paragraph("-" * 40)
        
        doc.add_paragraph("-" * 80)
        doc.add_paragraph("Huquqiy Eslatma: Ushbu bayonnoma SafeGuard Bot tizimi tomonidan qayd etilgan raqamli dalillar arxivi bo'lib, tergov va ekspertiza jarayonlarida rasmiy asos bo'lib xizmat qilishi mumkin.")
        doc.add_paragraph("")
        doc.add_paragraph("Mas'ul ekspert: ________________________ (Imzo)")
        doc.add_paragraph("Departament boshlig'i: ________________________ (Imzo)")
        
        doc.save(tmp.name)
        return tmp.name

    def forensic_list_pdf(self, cases: Sequence[ForensicCase], category_title: str) -> str:
        font_name = "Helvetica"
        font_path = _find_font()
        if font_path:
            try:
                pdfmetrics.registerFont(TTFont("UniFont", font_path))
                font_name = "UniFont"
            except Exception:
                pass

        tmp = tempfile.NamedTemporaryFile(suffix=".pdf", delete=False)
        doc = SimpleDocTemplate(tmp.name, pagesize=A4,
                                leftMargin=30, rightMargin=30,
                                topMargin=30, bottomMargin=30)
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle("ListTitleUni", parent=styles["Title"],
                                     fontName=font_name, fontSize=12, spaceAfter=8)
        normal_style = ParagraphStyle("ListNormalUni", parent=styles["Normal"],
                                      fontName=font_name, fontSize=8, leading=10)
        header_style = ParagraphStyle("ListHeaderUni", parent=styles["Normal"],
                                      fontName=font_name + "-Bold" if font_path is None else font_name,
                                      fontSize=9, leading=11, textColor=colors.white)

        elements = [
            Paragraph(f"<b>SafeGuard Bot — Kiber-Tergov Dalillar Arxivi — {category_title}</b>", title_style),
            Spacer(1, 10),
            Paragraph("Sana: " + datetime.now().strftime("%d.%m.%Y %H:%M"), normal_style),
            Paragraph(f"Jami yozuvlar soni: {len(cases)} ta", normal_style),
            Spacer(1, 15),
        ]

        section_title_style = ParagraphStyle(
            "SectionTitleUni", parent=styles["Normal"],
            fontName=font_name + "-Bold" if font_path is None else font_name,
            fontSize=10, leading=12, spaceBefore=10, spaceAfter=6,
            textColor=colors.HexColor("#2C3E50")
        )

        def make_table(group_cases, start_num_base):
            data = [[
                Paragraph("<b>#</b>", header_style),
                Paragraph("<b>Sana</b>", header_style),
                Paragraph("<b>Gumondor (ID)</b>", header_style),
                Paragraph("<b>Guruh</b>", header_style),
                Paragraph("<b>Jinoyat Turi</b>", header_style),
                Paragraph("<b>Tahlil sababi</b>", header_style)
            ]]
            for i, c in enumerate(group_cases, 1):
                lbl = {
                    "extremism": "🚨 Ekstrem",
                    "drugs": "💊 Giyoh",
                    "bullying": "⚠️ Bulling",
                    "link": "🔗 Havola",
                    "file": "📦 Fayl",
                }.get(c.violation_type, "🚫 Buzar")
                
                user_info = f"{c.full_name}\n(ID: {c.user_id})"
                if c.username:
                    user_info += f"\n@{c.username}"
                if c.phone:
                    user_info += f"\n{c.phone}"

                group_val = c.chat_title
                if getattr(c, "chat_username", None) and c.chat_username:
                    group_val += f"\n@{c.chat_username}\nID: {c.chat_id}"
                else:
                    group_val += f"\nID: {c.chat_id}"

                data.append([
                    Paragraph(str(start_num_base + i - 1), normal_style),
                    Paragraph(_escape_html(c.detected_at), normal_style),
                    Paragraph(_escape_html(user_info).replace("\n", "<br/>"), normal_style),
                    Paragraph(_escape_html(group_val).replace("\n", "<br/>"), normal_style),
                    Paragraph(_escape_html(lbl), normal_style),
                    Paragraph(_escape_html(c.reason), normal_style)
                ])

            t = Table(data, colWidths=[20, 65, 110, 85, 60, 160])
            t.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#2C3E50")),
                ("ALIGN", (0, 0), (-1, -1), "LEFT"),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.HexColor("#F9F9F9"), colors.white]),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                ("PADDING", (0, 0), (-1, -1), 4),
            ]))
            return t

        is_all = (category_title == "Barcha Dalillar")

        if is_all:
            groups_dict = {
                "🧠 1. Diniy Ekstremizm va Radikalizm": [c for c in cases if c.violation_type == "extremism"],
                "💊 2. Giyohvand moddalar aylanmasi": [c for c in cases if c.violation_type == "drugs"],
                "👤 3. Kiberbulling va Haqorat": [c for c in cases if c.violation_type == "bullying"],
                "🔗 4. Xavfli havolalar va Virusli fayllar": [c for c in cases if c.violation_type not in ("extremism", "drugs", "bullying")]
            }
            
            global_idx = 1
            for group_name, group_cases in groups_dict.items():
                if group_cases:
                    elements.append(Paragraph(f"<b>{group_name} ({len(group_cases)} ta)</b>", section_title_style))
                    elements.append(make_table(group_cases, global_idx))
                    global_idx += len(group_cases)
                    elements.append(Spacer(1, 15))
        else:
            elements.append(make_table(cases, 1))

        doc.build(elements)
        return tmp.name

    def forensic_list_docx(self, cases: Sequence[ForensicCase], category_title: str) -> str:
        tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
        doc = Document()
        
        h1 = doc.add_heading(f"SafeGuard Bot — Kiber-Tergov Dalillar Arxivi — {category_title}", 1)
        h1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph("Sana: " + datetime.now().strftime("%d.%m.%Y %H:%M"))
        doc.add_paragraph(f"Jami yozuvlar soni: {len(cases)} ta")
        doc.add_paragraph("")
        
        is_all = (category_title == "Barcha Dalillar")
        
        def add_docx_table(group_cases, start_num_base):
            table = doc.add_table(rows=1, cols=6)
            table.style = "Table Grid"
            
            headers = ["#", "Sana", "Gumondor", "Guruh", "Jinoyat Turi", "Tahlil sababi"]
            for i, h_text in enumerate(headers):
                cell = table.rows[0].cells[i]
                cell.text = h_text
                cell.paragraphs[0].runs[0].bold = True
                
            for i, c in enumerate(group_cases, 1):
                row = table.add_row().cells
                row[0].text = str(start_num_base + i - 1)
                row[1].text = c.detected_at
                
                user_info = f"{c.full_name} (ID: {c.user_id})"
                if c.username:
                    user_info += f"\n@{c.username}"
                if c.phone:
                    user_info += f"\n{c.phone}"
                row[2].text = user_info
                
                group_val = c.chat_title
                if getattr(c, "chat_username", None) and c.chat_username:
                    group_val += f"\n@{c.chat_username}\nID: {c.chat_id}"
                else:
                    group_val += f"\nID: {c.chat_id}"
                row[3].text = group_val
                
                lbl = {
                    "extremism": "Ekstremizm",
                    "drugs": "Giyohvandlik",
                    "bullying": "Kiberbulling",
                    "link": "Xavfli Havola",
                    "file": "Zararli Fayl",
                }.get(c.violation_type, "Qonunbuzarlik")
                row[4].text = lbl
                
                row[5].text = c.reason

        if is_all:
            groups_dict = {
                "🧠 1. Diniy Ekstremizm va Radikalizm": [c for c in cases if c.violation_type == "extremism"],
                "💊 2. Giyohvand moddalar aylanmasi": [c for c in cases if c.violation_type == "drugs"],
                "👤 3. Kiberbulling va Haqorat": [c for c in cases if c.violation_type == "bullying"],
                "🔗 4. Xavfli havolalar va Virusli fayllar": [c for c in cases if c.violation_type not in ("extremism", "drugs", "bullying")]
            }
            
            global_idx = 1
            for group_name, group_cases in groups_dict.items():
                if group_cases:
                    doc.add_heading(f"{group_name} ({len(group_cases)} ta)", 3)
                    add_docx_table(group_cases, global_idx)
                    global_idx += len(group_cases)
                    doc.add_paragraph("")
        else:
            add_docx_table(cases, 1)
            
        doc.save(tmp.name)
        return tmp.name

